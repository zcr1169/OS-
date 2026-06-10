#!/usr/bin/env python3
"""生成操作系统A课程设计实习报告——完整版"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

RGB_BLACK = RGBColor(0,0,0)

def set_run_font(run, name='宋体', size=12, bold=False, color=RGB_BLACK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color

def add_para(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    set_run_font(run, '宋体', size, bold)
    return p

def add_body(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, '宋体', 12, bold)
    return p

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, '黑体', 16 if level==1 else 14, bold=True)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_sub(text):
    """缩进子段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, '宋体', 12)
    return p

def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('• ' + text)
    set_run_font(run, '宋体', 12)
    return p

# ==================== 封面 ====================
add_para('北 京 林 业 大 学', size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))
add_para('2025—2026学年第2学期  操作系统A  实习报告书', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(30))

info_items = [
    ('专    业：', '计算机科学与技术', '班    级：', '信息24-2'),
    ('姓    名：', '周宸冉', '学    号：', '241001209'),
    ('实习地点：', '北京林业大学', '辅导教师：', '孟伟'),
]
for left_label, left_val, right_label, right_val in info_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{left_label}  {left_val}')
    set_run_font(run, '宋体', 12)
    run = p.add_run(f'    {right_label}  {right_val}')
    set_run_font(run, '宋体', 12)

doc.add_paragraph()
add_body('实习内容：  可持久化的操作系统核心模拟器设计与实现')
add_body('实习环境：  Windows 11 + MSVC (C++17)，命令行交互式终端')
doc.add_page_break()

# ==================== 正文 ====================

# ════════════════════════════════════════
add_heading_styled('一、实习步骤', level=1)
add_body('本次课程设计共32学时，按四个阶段推进，具体安排如下：')
add_body('第一阶段（第1-8学时）——需求分析与设计：深入理解课设目标，明确系统功能需求。完成核心数据结构定义（PCB结构体、调度队列、内存块链表等），确定二进制持久化文件格式，绘制系统架构图和主要流程图。本阶段的重点是确认所有设计细节，避免编码阶段返工。')
add_body('第二阶段（第9-20学时）——编码实现：依据设计文档完成全部代码编写。具体包括用户管理模块（注册/登录/密码锁定）、进程管理模块（10个核心命令）、MLFQ多级反馈队列调度器、动态分区内存管理（三种分配算法）、二进制持久化（save/load）、多线程架构（前台线程+后台线程+调度线程）、多实例共享机制（文件锁+观察者轮询）和overview可视化命令。')
add_body('第三阶段（第21-28学时）——调试与验证：重点攻克多线程环境下的并发问题与死锁风险。模拟多种边界场景，包括：空队列调度、非法参数输入、高负载多进程并发、双窗口多实例同步等。进行全流程黑盒测试，验证持久化save/load的完整性和重启恢复能力。')
add_body('第四阶段（第29-32学时）——报告与验收：整理开发过程中的技术文档，撰写包含设计思路、核心数据结构和算法实现、测试结果的课程设计报告。准备现场演示环境，完成最终验收。')

# ════════════════════════════════════════
add_heading_styled('二、技术方案', level=1)

add_heading_styled('2.1 系统总体架构', level=2)
add_body('本系统采用C++17开发，MSVC编译器，Windows平台命令行交互。源码共7个.cpp文件和7个.h头文件，约3500行代码。整体架构分为三个层次：')
add_body('（1）用户交互层：前台主循环读取用户输入，通过命令解析器（CommandParser）将字符串解析为命令名和参数列表，推入线程安全消息队列。')
add_body('（2）业务逻辑层：后台工作线程从消息队列取出命令并执行，操作PCB表、内存表和用户表。调度器线程独立运行，每2秒执行一次MLFQ调度。')
add_body('（3）数据持久层：每执行完一条命令后，系统将全量状态序列化到二进制文件os_state.bin。观察者实例通过轮询文件时间戳同步数据。')
add_body('三个线程通过消息队列（mutex + condition_variable）解耦通信，共享数据结构通过recursive_mutex保护，统一锁顺序（PM锁→Sched锁）防止死锁。')

add_heading_styled('2.2 用户管理模块', level=2)
add_body('用户信息存储在unordered_map<string, UserInfo>中，UserInfo包含用户名、FNV-1a密码哈希值、连续失败次数和锁定标志。register命令检查用户名冲突后哈希存储密码；login命令比对哈希值，连续3次错误锁定账户；logout命令清空当前登录用户。所有命令按当前登录用户过滤进程列表，实现用户隔离。')

add_heading_styled('2.3 进程管理模块', level=2)
add_body('进程控制块PCB是核心数据结构，包含pid、ppid、name、state（6种状态枚举）、priority（0-15）、cpuTime、burstTime、totalMemory、children列表、owner和memoryBlocks列表。')
add_body('进程管理器ProcessManager维护一个unordered_map<int32_t, PCB>哈希表，以PID为键。共实现10个命令：create_pcb（分配PID并创建PCB，默认入Q0）、kill_pcb（递归级联删除子进程后再删除自身，回调释放调度队列和内存）、block_pcb（状态改为BLOCKED并从调度队列移除）、wakeup_pcb（状态改回READY并重置Q0）、show_pcb（格式化输出PCB全部字段）、list_pcb（列表显示当前用户进程，用padRight中文对齐）、ptree（递归遍历children画树形结构）、suspend（移出调度队列，保留内存）、resume（重新入Q0）、renice（修改优先级后dequeue+enqueue移到新队列）。')

add_heading_styled('2.4 MLFQ调度器', level=2)
add_body('调度器实现三级多级反馈队列：Q0（优先级0-3，时间片2）、Q1（优先级4-7，时间片4）、Q2（优先级8-15，时间片8）。核心设计规则包括：')
add_body('（1）入队规则：新进程统一入Q0，时间片耗尽降级到下一级，阻塞唤醒后重置Q0，恢复挂起也重置Q0，renice按新优先级入对应队列。')
add_body('（2）调度规则：严格按Q0→Q1→Q2优先级扫描，高优先级队列为空时才调度低优先级。step命令输出完整决策链路：当前队列快照→扫描过程→选中进程→时间片分配→CPU时间推进→完成/降级判断→执行后队列状态。')
add_body('（3）老化机制：每调度10次触发aging，将Q2队首提升到Q1（优先级改7），Q1队首提升到Q0（优先级改3），防止低优先级进程饥饿。')
add_body('（4）自动调度：独立schedulerLoop线程每2秒调用一次step，连续3次发现队列全空则自动停止，避免空转。')
add_body('（5）进程完成：当cpuTime ≥ burstTime时，进程自动终止，回调释放所有内存块，输出完成日志。')

add_heading_styled('2.5 内存管理模块', level=2)
add_body('模拟1024KB动态分区内存分配，维护两个链表：freeBlocks_（空闲块链表，按地址有序）和allocBlocks_（已分配块链表）。')
add_body('（1）分配算法：首次适应（FF）选第一个够大的空闲块，适合快速分配；最佳适应（BF）选最接近请求大小的块，减少内部碎片；最坏适应（WF）选最大的块，避免产生小碎片。三种算法通过switch切换。')
add_body('（2）非进程目标：alloc命令支持分配给进程（指定PID），也支持分配给数据区（data）、IO区（io）、内核区（kernel），分别用特殊PID常量-2、-3、-4标识，show_mem时显示文字标注。')
add_body('（3）内存释放：freeByPid释放指定进程的所有内存块并自动合并相邻空闲块；freeByAddr按地址释放单块。')
add_body('（4）碎片管理：compact命令合并相邻空闲块消除外部碎片。碎片率计算公式为：1 -（最大连续空闲块/总空闲空间）× 100%。')
add_body('（5）换出换入：swap_out将进程内存块记录保存到swappedOut_映射表后释放物理内存；swap_in重新分配内存恢复进程内存。')

add_heading_styled('2.6 持久化模块', level=2)
add_body('采用二进制格式保存全量系统状态。文件格式依次为：魔数（0x4F535353）+版本号+时间戳，之后依次为用户数据（用户名、密码哈希、失败次数、锁定状态、当前登录用户）、进程数据（每个PCB的全部字段，含children列表和memoryBlocks）、内存数据（空闲块链表、已分配块链表、分配算法枚举）、调度器数据（三个队列的PID顺序、运行标志）。')
add_body('save时逐模块加锁后序列化写入二进制文件。load时先读取到临时变量并校验魔数和版本，校验通过后暂停调度器、清空现有状态、原子替换全部数据、恢复调度器。这种设计防止读取到一半文件损坏时破坏已有状态。程序启动时自动检测文件并调用load，实现重启无缝恢复。')

add_heading_styled('2.7 多线程与多实例', level=2)
add_body('系统采用三线程架构：前台交互线程（main循环，负责读取用户输入→解析→推入消息队列→等待结果输出）、后台工作线程（backendLoop，从消息队列取命令→executeCommand→结果推入resultQueue）、调度器线程（schedulerLoop，每2秒step一次，独立运行）。')
add_body('线程间通过MessageQueue<T>模板类通信，基于std::mutex和std::condition_variable实现。前台push后notify_one唤醒后台pop。每个管理器自带recursive_mutex，统一锁顺序为PM锁→Sched锁。')
add_body('多实例共享采用文件锁方案：先启动的实例通过CreateFileW+LockFileEx获取文件锁成为后端，后启动的成为观察者。观察者将命令写入commands.txt，后端定时检查并执行。观察者同时每200ms轮询os_state.bin的时间戳变化自动load同步。后端退出后文件锁自动释放，观察者下一次tryLock成功则升级为新的后端。')

add_heading_styled('2.8 可视化模块', level=2)
add_body('overview命令在58列宽的字符画框内展示系统全景快照，包含四个板块：Process Tree（进程树，按用户过滤）、Memory Map（ASCII字符画内存分布图，用|##PID(sizeKB)|--free(sizeKB)--|格式）、多级反馈队列（三队列进程列表，显示名称(PID)[cpuTime/burstTime]）、Stats（就绪/运行/阻塞/挂起计数及调度器状态和内存算法）。')

# ════════════════════════════════════════
add_heading_styled('三、设计思想', level=1)
add_body('本课程设计的目标是设计并实现一个模拟操作系统核心功能的交互式环境，让抽象的内核原理可视化、可交互。核心设计思想包括以下四点：')
add_body('1. 模块解耦：将系统拆分为进程管理、内存管理、用户管理、调度器、持久化五个独立模块，每个模块有明确的接口和职责，通过OSSimulator主控类串联。这样各模块可独立开发、测试和维护。')
add_body('2. 消息驱动架构：前台与后台通过消息队列通信，前台只管读输入和输出结果，后台负责实际执行。这种生产者消费者模式有效解耦了UI操作与内核计算，前台不会因为后台耗时操作而卡顿。')
add_body('3. 多实例共享方案：采用文件锁+文件轮询的轻量级方案实现多终端并发访问同一内核状态。相比命名管道或共享内存，文件锁方案实现简单、跨平台、进程退出自动释放无残留。')
add_body('4. MLFQ调度算法选择：多级反馈队列是操作系统教材中的经典调度算法，能兼顾短进程的快速响应和长进程的吞吐量。本实现通过三级队列、时间片倍增、降级和老化机制，完整呈现了MLFQ的核心特性。')

# ════════════════════════════════════════
add_heading_styled('四、流程图', level=1)
add_body('（注：以下为文字描述流程图，最终报告中应替换为Visio或绘图工具绘制的流程图截图）')

add_heading_styled('4.1 系统总体架构图', level=2)
add_body('用户输入 → 命令解析器 → (前端线程) → cmdQueue消息队列 → (后端线程) → 命令分发器')
add_body('                                                                        ↓')
add_body('                                                          ┌──── 进程管理器 ────┐')
add_body('                                                          │   内存管理器       │')
add_body('                                                          │   用户管理器       │')
add_body('                                                          │   调度器           │')
add_body('                                                          │   持久化           │')
add_body('                                                          └──────────────────┘')
add_body('                                                                        ↑')
add_body('                                                        (调度器线程独立运行)')

add_heading_styled('4.2 MLFQ调度流程图', level=2)
add_body('新进程创建 → enqueue(pid, 0) → 入Q0')
add_body('    ↓')
add_body('step() 扫描 Q0→Q1→Q2')
add_body('    ↓')
add_body('找到第一个非空队列，取队首PID')
add_body('    ↓')
add_body('state = RUNNING, cpuTime += 时间片')
add_body('    ↓')
add_body('判断：cpuTime ≥ burstTime?')
add_body('    ├── 是 → state = TERMINATED，释放内存，输出"进程完成"')
add_body('    └── 否 → state = READY，降级到下一队列（Q0→Q1→Q2→Q2）')
add_body('                ↓')
add_body('          scheduleCount++')
add_body('                ↓')
add_body('         scheduleCount % 10 == 0?')
add_body('                ├── 是 → agePriorities(): Q2→Q1, Q1→Q0')
add_body('                └── 否 → 继续')

add_heading_styled('4.3 内存分配流程图', level=2)
add_body('alloc(size, pid)')
add_body('    ↓')
add_body('switch(选算法):')
add_body('    ├── FIRST_FIT：遍历freeBlocks，第一个>=size即选')
add_body('    ├── BEST_FIT：遍历全部，选size差距最小的')
add_body('    └── WORST_FIT：遍历全部，选size差距最大的')
add_body('    ↓')
add_body('选中的空闲块.size == size?')
add_body('    ├── 是 → 整个块从freeBlocks移到allocBlocks')
add_body('    └── 否 → 分裂：前部分分配，后部分继续空闲')
add_body('    ↓')
add_body('更新PCB的memoryBlocks和totalMemory')
add_body('    ↓')
add_body('返回分配起始地址')

# ════════════════════════════════════════
add_heading_styled('五、主要数据结构说明', level=1)

add_heading_styled('5.1 PCB（进程控制块）', level=2)
add_body('PCB是系统中最重要的数据结构，每个进程对应一个PCB。定义在进程控制块.h中，包含：')
add_bullet('pid：进程唯一标识符，从1开始自增分配')
add_bullet('ppid：父进程PID，-1表示无父进程（init的ppid=-1）')
add_bullet('name：进程名称，由用户在create_pcb时指定')
add_bullet('state：状态枚举，共6种——NEW(新建)、READY(就绪)、RUNNING(运行中)、BLOCKED(阻塞)、SUSPENDED(挂起)、TERMINATED(终止)')
add_bullet('priority：调度优先级，范围0-15，数值越小优先级越高')
add_bullet('cpuTime：已消耗的CPU时间片累计值，每次step时增加')
add_bullet('burstTime：进程需要的总CPU时间，cpuTime达到此值后进程自动终止')
add_bullet('totalMemory：进程当前占用的总内存大小，单位KB')
add_bullet('children：子进程PID的vector，用于ptree遍历和级联删除')
add_bullet('owner：所属用户名，用于list_pcb/ptree/overview时的用户隔离过滤')
add_bullet('memoryBlocks：已分配内存块列表，每个元素为(起始地址, 大小)对')

add_heading_styled('5.2 调度队列结构', level=2)
add_body('调度器维护一个vector<deque<int32_t>>，包含三个双端队列Q0、Q1、Q2。每个队列存储的是进程的PID（而非PCB指针），这样即使PCB被删除也不会出现悬垂指针。入队时根据priority参数计算队列索引（0→Q0，4→Q1，8→Q2），出队时遍历三个队列查找并移除。deque支持从两端操作，队首出队、队尾入队，符合FIFO调度需求。')

add_heading_styled('5.3 内存块结构', level=2)
add_body('MemBlock结构体包含四个字段：startAddr（起始地址KB）、size（大小KB）、pid（所属进程PID，-1表示空闲，-2数据，-3IO，-4内核）、free（是否空闲）。内存管理器维护两个list<MemBlock>：freeBlocks_按地址排序存储空闲块，allocBlocks_存储已分配块。采用list容器因为分配释放时需要频繁插入删除，list的插入删除操作为O(1)。')

add_heading_styled('5.4 核心函数说明', level=2)
add_body('ProcessManager::createPCB()：分配PID、构造PCB、插入哈希表、更新父进程children。返回新PID，父进程不存在时返回-1。')
add_body('ProcessManager::killPCB()：从父进程children移除自身→递归杀子进程→回调释放调度队列和内存→清空PCB内存记录→从哈希表删除。ptree依赖的children关系也一并清理。')
add_body('Scheduler::step()：核心调度函数。输出当前三队列快照→按Q0→Q1→Q2扫描→取队首→RUNNING+cpuTime累加→判断完成或降级→输出执行后队列。每步输出完整决策链路，满足验收要求。')
add_body('MemoryManager::alloc()：按FF/BF/WF算法从freeBlocks_选择空闲块，分裂并分配。返回起始地址，空间不足返回-1。')
add_body('MemoryManager::compact()：sort空闲块按地址排序，遍历合并相邻块（it->startAddr + it->size == next->startAddr）。')
add_body('StateSerializer::save()：逐模块加锁→writeVal写入二进制→writeStr写入字符串（先写int32长度再写内容）。魔数0x4F535353用于文件格式校验。')
add_body('StateSerializer::load()：读魔数校验→读版本校验→读入临时变量→暂停调度器→清空旧状态→原子替换全部数据→恢复调度器。队列中的死PID自动清理。')
add_body('OSSimulator::run()：前台主循环。打印提示符→读输入→解析命令→exit/help直接处理→后端模式推cmdQueue等resultQueue→观察者模式写commands.txt+轮询认证结果。')

# ════════════════════════════════════════
add_heading_styled('六、实验结果', level=1)
add_body('以下为各模块的测试运行结果：')

add_heading_styled('6.1 用户管理', level=2)
add_body('注册、重复注册、登录、错误密码提示、登出后命令拦截等功能均正常。连续输错3次密码后账户锁定，再次登录提示"账户已被锁定"。用户隔离验证：用户A创建进程后，用户B登录执行list_pcb看不到A的进程。')

add_heading_styled('6.2 进程管理', level=2)
add_body('10个进程命令全部通过测试。create_pcb成功创建进程并返回PID；kill_pcb级联删除子进程，父进程和子进程均从进程列表消失；ptree正确显示进程树形结构（含缩进和├─符号）；block_pcb状态变为阻塞；wakeup_pcb恢复就绪；suspend移出调度队列；resume恢复并重置Q0；renice修改优先级后队列位置立即更新。')

add_heading_styled('6.3 MLFQ调度器', level=2)
add_body('自动调度测试：创建3个CPU密集型进程（burstTime=10），执行start_sched后约10-15秒全部运行完毕。日志显示进程按Q0→Q1→Q2逐级降级，时间片2→4→8递增，CPU时间动态增长非静态修改。单步调度测试：每个step输出完整的"当前队列→选中进程→时间片→CPU时间→完成/降级→执行后队列"链路。老化机制在10次调度后自动触发。')

add_heading_styled('6.4 内存管理', level=2)
add_body('alloc/free_mem/show_mem/compact/mem_stat/set_alloc_algo/pgfault/swap_out全部通过。碎片率计算正确（碎片率=1-最大连续空闲/总空闲）。compact后相邻空闲块合并，碎片率降为0。三种分配算法切换后alloc行为不同。swap_out释放进程内存，swap_in成功恢复。非进程目标alloc data/io/kernel在show_mem中显示文字标注。')

add_heading_styled('6.5 持久化', level=2)
add_body('save命令保存全量状态到二进制文件。退出程序后重新启动，系统自动加载os_state.bin，进程、内存、调度队列全部精确恢复。list_pcb和show_mem确认进程状态和内存分配与保存前完全一致。队列顺序和已执行CPU时间无损。')

add_heading_styled('6.6 全景可视化', level=2)
add_body('overview命令正确显示进程树、内存ASCII分布图、MLFQ三队列状态和快速统计信息。执行进程创建/调度/内存操作后，overview内容同步更新，非静态输出。')

# ════════════════════════════════════════
add_heading_styled('七、实验分析及感受', level=1)

add_heading_styled('7.1 进程管理方面', level=2)
add_body('通过亲手实现PCB的6种状态转换和10个操作命令，对进程的生命周期管理有了直观认识。级联杀子进程的实现让我深入理解了父子进程关系树的管理复杂度——杀掉父进程时必须递归清理所有子孙，否则会产生孤儿进程。ptree的递归遍历算法虽然简单，但要处理好缩进层级和孤儿进程的显示并不容易。')

add_heading_styled('7.2 调度算法方面', level=2)
add_body('MLFQ调度器的实现让我深刻理解了多级反馈队列的设计哲学：短进程在Q0快速完成获得良好响应，长进程逐步降级到Q2获得长时间片减少切换开销。老化机制的实现验证了"防止饥饿"这一调度设计原则的实际意义——没有aging的话，Q2的进程可能永远得不到CPU。调试step的完整决策链路输出时，我意识到好的调度器不仅算法要正确，日志输出也要清晰可读，否则老师无法验证调度器的动态行为。')

add_heading_styled('7.3 内存管理方面', level=2)
add_body('三种分配算法的对比实现让我理解了不同算法的适用场景：FF最快但可能产生较多碎片，BF能减少碎片但遍历较慢，WF倾向于保留中等大小的空闲块。内存碎片管理是实际操作系统中非常重要的问题，compact操作展示了解决外部碎片的典型方法。非进程目标（data/io/kernel）的支持让我了解了操作系统内存不仅分配给进程，还要分配给内核模块和IO缓冲区。')

add_heading_styled('7.4 多线程并发方面', level=2)
add_body('消息队列+互斥锁的方案让我体会了生产者消费者模型的实用性。调试多线程竞态条件是最耗时的部分——通过统一锁顺序（PM锁→Sched锁）解决了潜在的死锁问题。recursive_mutex的使用简化了递归函数的加锁逻辑。多实例共享中文件锁+轮询的方案虽然是简化实现，但涵盖了分布式系统中常见的锁竞争和状态同步问题。')

add_heading_styled('7.5 遇到的困难与解决', level=2)
add_body('（1）死锁问题：调度器step()需要同时持有PM锁和Sched锁，如果两个线程以不同顺序加锁就会死锁。解决方案：统一加锁顺序，所有地方先加PM锁再加Sched锁。')
add_body('（2）数据目录路径问题：初始使用相对路径"./数据/"，从不同目录启动程序时数据落盘位置不同。解决方案：用GetModuleFileNameW获取exe路径，数据文件统一放在exe旁边。')
add_body('（3）持久化安全性：load时如果文件损坏可能导致系统状态丢失。解决方案：先读入临时变量并校验魔数和版本，校验通过后再替换现有状态。')
add_body('（4）观察者登录验证：初始实现中观察者端login不验证密码直接标记为已登录。解决方案：写入commands.txt后同步等待后端执行结果，检查userMgr_.isLoggedIn()确认登录是否成功。')
add_body('（5）进程自动终止内存泄漏：进程跑完burstTime后自动终止，但占用的内存未释放。解决方案：在step()中判断终止时回调onTerminate_释放内存。')

add_heading_styled('7.6 总结', level=2)
add_body('这次课程设计让我从"操作系统使用者"变成了"操作系统实现者"。在编写代码之前，我对进程调度、内存分配、文件系统的理解仅停留在理论层面——知道MLFQ、知道首次适应算法、知道PV操作，但从未亲手实现过。通过这次实践，我不仅深入理解了每个算法的工作原理，还体会到理论到实现之间的距离：教材上的算法描述可能只有几行伪代码，但真正实现需要考虑边界条件、并发安全、错误处理和用户体验等诸多因素。同时，调试多线程程序和修复各种bug的过程也锻炼了我的问题排查能力。总的来说，这是一次非常有价值的实践经历，让我对操作系统这门课程有了更深刻、更立体的认识。')

# ===== 保存 =====
out_path = os.path.join(os.path.dirname(__file__), '课设报告', '实习报告.docx')
doc.save(out_path)
print(f'报告已生成: {out_path}')
